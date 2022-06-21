"""
@author:maohui
@time:2022/6/20 10:12
  　　　　　　　 ┏┓    ┏┓+ +
  　　　　　　　┏┛┻━━━━┛┻┓ + +
  　　　　　　　┃        ┃ 　 
  　　　　　　　┃     ━  ┃ ++ + + +
  　　　　　 　████━████ ┃+
  　　　　　　　┃        ┃ +
  　　　　　　　┃   ┻    ┃
  　　　　　　　┃        ┃ + +
  　　　　　　　┗━┓   ┏━━┛
  　　　　　　　  ┃   ┃
  　　　　　　　  ┃   ┃ + + + +
  　　　　　　　  ┃   ┃　　　Code is far away from bug with the animal protecting
  　　　　　　　  ┃   ┃+ 　　　　神兽保佑,代码无bug
  　　　　　　　  ┃   ┃
  　　　　　　　  ┃   ┃　　+
  　　　　　　　  ┃   ┗━━━━━━━┓ + +     
  　　　　　　　  ┃           ┣┓
  　　　　　　　  ┃           ┏┛
  　　　　　　　  ┗┓┓┏━━━━━┳┓┏┛ + + + +
  　　　　　　　   ┃┫┫     ┃┫┫
  　　　　　　　   ┗┻┛     ┗┻┛+ + + +
"""
"""
文档（document）
段落（paragraph）
行内元素（runs）    内容（text）    表格（tables）
内容（text）                      行（table_rows）     列(table_columns)
字体（font）                      单元格(table_cell)   单元格(table_cell) 
颜色（color）                     内容（text）          内容（text）
字号（size）
"""
from docx import Document
from docx.enum.style import WD_STYLE_TYPE
import re


# 添加标题
# docx1.add_heading('我是一级标题',level=1)
# 加分页符
# docx1.add_page_break()
# 加正文
# docx1.add_paragraph("我是正文")
# 加文字块
# a=docx1.add_paragraph("我是正文在我后面添加的文字会被设置格式") #加正文
# a.add_run("加粗").bold=True
# a.add_run("普通")
# a.add_run("斜体").italic=True
# 指定第二个段落前插入一个新段落
# para=docx1.paragraphs[1]
# para.insert_paragraph_before("这是新添加的第二段")
# #保存
# docx1.save("C:/Users/PC/Desktop/web_2_requires.docx")


def set_picture(docx1):
    """"按图片比例和百分比设置图片"""
    pict1 = docx1.paragraphs[1].add_run().add_picture("...1")  # 在文档第二个段落插入图片
    pict2 = docx1.paragraphs[1].add_run().add_picture("...2")  # 在文档第三个段落插入图片
    print("打印第一个图片的高和第二个图片的高", docx1.inline_shapes[0].height, docx1.inline_shapes[0].height)
    # 打印第2个图片的高和第1个图片的高度的比值
    height_rate = docx1.inline_shapes[1].height / docx1.inline_shapes[0].height
    # 打印第2个图片的高和第1个图片的宽度的比值
    width_rate = docx1.inline_shapes[1].width / docx1.inline_shapes[0].width
    # 安装比例设置图片1高度
    pict1.height = int(docx1.inline_shapes[0].height * height_rate)
    # 安装比例设置图片1宽度
    pict1.height = int(docx1.inline_shapes[0].height * width_rate)
    # 按照50%比例设置图片高度
    pict2.height = int(docx1.inline_shapes[1].height * 0.5)


def text_font_adjust(docx1):
    """1.1建议撰写内容（宋体，五号，加粗）
        1.1.1具体要求题目（宋体，五号，加粗）10.5
        （1）子题目（宋体，五号，加粗）"""
    from docx.shared import Pt, RGBColor  # 字号，颜色
    from docx.oxml.ns import qn  # 中文字体
    for para in docx1.paragraphs:
        # 标题都加粗
        # if re.match("^Heading \d+$",para.style.name):
        if para.style.name == 'Heading 3':
            for block in para.runs:
                block.font.bold = True
                block.font.italic = False
        if para.style.name == 'Heading 4':
            for block in para.runs:
                block.font.bold = True
                block.font.italic = False
        # if para.style.name=='List Paragraph':

        # 所有的中英文都是10.5磅
        for block in para.runs:
            # block.font.bold = True
            # block.font.italic = True
            # 设置数字和英文(西文)的样式
            block.font.name = '宋体'
            # 设置汉字（中文）的样式
            block._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
            block.font.size = Pt(10.5)


def para_retract(docx1):
    """段落要求：段前、段后0行，行距1.5倍"""
    from docx.shared import Cm, Pt  # pt-磅（int）、cm-厘米（float）、inches-英寸（float）、mm-毫米（float）
    from docx.enum.text import WD_LINE_SPACING
    from docx.shared import Inches
    for para in docx1.paragraphs:
        # 1.5倍行距
        para.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
        # 段前
        para.paragraph_format.space_before = Pt(0)
        # 段后
        para.paragraph_format.space_after = Pt(0)
        if para.style.name=='Heading 5':
        # 左缩进
            para.paragraph_format.left_indent=Cm(0)
            #首行缩进两个字符
            para.style.font.size = Pt(10.5)
            para.paragraph_format.first_line_indent = para.style.font.size * 2
        # 只修改正文i.次一级目录或内容（宋体，五号，编号对齐方式：左侧对齐；段落左侧1cm，悬挂0.75cm）
        # if para.style.name=='List Paragraph':
        #     # 首行缩进
        #     # para.paragraph_format.first_line_indent = Inches(0.3)
        #     # 悬挂缩进，除了首行都缩进
        #     para.paragraph_format.first_line_indent=Cm(-0.75)


def table_set(docx1):
    """表格文字：宋体，五号，行距：单倍行距"""
    from docx.shared import Cm, Pt  # pt-磅（int）、cm-厘米（float）、inches-英寸（float）、mm-毫米（float）
    from docx.oxml.ns import qn  # 中文字体
    # 循环所有表格
    for index in range(len(docx1.tables)):
        # 循环表格的行
        for row in docx1.tables[index].rows:
            # 循环单元格
            for cell in row.cells:
                # 循环单元格的段落
                for para in cell.paragraphs:
                    # 单倍行间距
                    para.paragraph_format.line_spacing = 1.0
                    # 循环文本块
                    for block in para.runs:
                        block.font.name = '宋体'
                        block.font.size = Pt(10.5)
                        block._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')


def check_title_styles_name(docx1):
    """# 查看所有标题的类型"""
    titles = docx1.styles
    for i in titles:
        if i.type == WD_STYLE_TYPE.PARAGRAPH:
            print(i.name)
# count=0
# for para in docx1.paragraphs:#遍历段落
#     #标题筛选
#     # if re.match("^Heading \d+$",para.style.name):
#     if para.style.name=='Heading 1':
#         print(para.text)
#     #寻找一共有几个“医学合作”
#     if '医学合作' in para.text:
#         count += 1
#     block=para.runs#每一段落以文字格式分成块
#     for text1 in block:
#         print(text1.text)
# print(count)

if __name__ == "__main__":
    # 打开
    docx1 = Document("C:/Users/PC/Desktop/5.3明确产品主要技术指标及确定依据0615.docx")
    # for block in para.runs:
    #     if re.match("^\d+.*:$",block.text)==True:
    #         block.font.bold=True
    text_font_adjust(docx1)
    para_retract(docx1)
    table_set(docx1)
    # 保存
    docx1.save("C:/Users/PC/Desktop/5.3明确产品主要技术指标及确定依据0615.docx")
